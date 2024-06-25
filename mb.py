import docx
from decimal import Decimal

def mb(type, max):
    """
    生成并保存指定类型的表格文档（Π表、平方表、乘法表）。

    参数:
    - type: str，指定生成的表格类型，可选值为"Π表"、"平方表"、"乘法表"。
    - max: int，指定表格的上限值，生成的表格将包含从1到max的数值。

    返回值:
    - 无。函数将生成的文档保存为[type].docx格式。
    """
    doc = docx.Document()  # 创建一个新的docx文档

    if type == "Π表":
        doc.add_heading("Π表")  # 添加表头
        pa = doc.add_paragraph()  # 创建一个段落
        pi = Decimal(3.14)  # 定义π的值
        # 循环添加Π表的内容
        for i in range(1, max + 1):
            id = Decimal(i)
            pa.add_run(str(i)+"Π≈"+str(id*pi.quantize(Decimal('0.01')))+"    ")

    if type == "平方表":
        doc.add_heading("平方表")  # 添加表头
        pa = doc.add_paragraph()  # 创建一个段落
        # 循环添加平方表的内容
        for i in range(1, max + 1):
            id = Decimal(i)
            pa.add_run(str(i)+"²="+str(id*id)+"    ")

    if type == "乘法表":
        doc.add_heading("乘法表")  # 添加表头
        pa = doc.add_paragraph()  # 创建一个段落
        # 双重循环添加乘法表的内容
        for i in range(1, max + 1):
            for j in range(1, max + 1):
                id = Decimal(i)
                jd = Decimal(j)
                pa.add_run(str(i)+"×"+str(j)+"="+str(id*jd)+"    ")

    doc.save(str(type)+".docx")  # 保存文档

